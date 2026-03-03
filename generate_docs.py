#!/usr/bin/env python3
"""Generate Word documents for RestartiX offer and implementation guide."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = "/Users/david-adrianbabtan/Desktop/website-uri/restartix"

# ═══════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows, header_color="2F5496"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(cell, header_color)
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return h

def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.style = doc.styles['List Bullet 2'] if 'List Bullet 2' in [s.name for s in doc.styles] else doc.styles['List Bullet']
        p.paragraph_format.left_indent = Cm(1.5 * level)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


# ═══════════════════════════════════════════════════════════
# DOCUMENT 1: OFERTA
# ═══════════════════════════════════════════════════════════

def create_oferta():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Title ──
    title = doc.add_heading('Ofertă Optimizare WooCommerce', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('RestartiX – magazinrestartix.ro & zerodurere.net')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_para(doc, '')
    
    p = doc.add_paragraph()
    run = p.add_run('Propunere one-time | Full Out Media')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    p2 = doc.add_paragraph()
    run = p2.add_run('Data: Martie 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    doc.add_page_break()
    
    # ── Context ──
    add_heading_styled(doc, '1. Context & Obiective', level=1)
    
    add_para(doc, 'Pe baza discuțiilor avute și a informațiilor primite despre ecosistemul digital RestartiX, prezentăm mai jos oferta noastră one-time pentru optimizarea celor două magazine WooCommerce.')
    
    add_para(doc, 'Probleme identificate:', bold=True)
    add_bullet(doc, 'Viteză redusă a site-urilor, în special la peak-uri de trafic (100-150 utilizatori concurenți)')
    add_bullet(doc, 'Erori de accesibilitate în timpul evenimentelor live')
    add_bullet(doc, 'Lipsa fluxurilor automate de comunicare e-commerce (abandon cart, tracking AWB, follow-up)')
    add_bullet(doc, 'Checkout neoptimizat pentru conversie')
    add_bullet(doc, 'Lipsa comunicării automate pe WhatsApp')
    
    doc.add_page_break()
    
    # ── Pachet 1 ──
    add_heading_styled(doc, '2. Pachet 1 – Audit & Optimizare Performanță Tehnică', level=1)
    add_para(doc, 'Ambele magazine: magazinrestartix.ro + zerodurere.net', bold=True)
    
    add_heading_styled(doc, '2.1 Audit complet de performanță', level=2)
    add_bullet(doc, 'Analiză Core Web Vitals (LCP, FID, CLS) pe ambele magazine')
    add_bullet(doc, 'Identificare bottleneck-uri: pluginuri grele, query-uri lente, resurse render-blocking')
    add_bullet(doc, 'Audit hosting: configurare PHP, MySQL, limits server (VPS custom + Romarg WordPress Pro-4)')
    add_bullet(doc, 'Raport detaliat cu findings + recomandări prioritizate')
    
    add_heading_styled(doc, '2.2 Optimizare viteză', level=2)
    add_bullet(doc, 'Configurare cache avansat (Redis/Memcached object cache pe VPS, page cache)')
    add_bullet(doc, 'Optimizare imagini (conversie WebP, lazy loading, CDN)')
    add_bullet(doc, 'Minificare & combinare CSS/JS, eliminare resurse render-blocking')
    add_bullet(doc, 'Optimizare bază de date WooCommerce (cleanup revisions, transients, expired sessions)')
    add_bullet(doc, 'Preload fonturi critice, defer scripturi non-critice')
    
    add_heading_styled(doc, '2.3 Stabilitate la peak traffic', level=2)
    add_bullet(doc, 'Configurare corectă VPS (magazinrestartix.ro): PHP workers, OPcache, MySQL tuning')
    add_bullet(doc, 'Implementare CDN (Cloudflare/BunnyCDN) pe ambele magazine')
    add_bullet(doc, 'Rate limiting + protecție bot traffic în timpul evenimentelor')
    add_bullet(doc, 'Testare de stres post-optimizare cu simulare trafic de peak')
    add_bullet(doc, 'Raport before/after cu metrici concrete (TTFB, LCP, load time)')
    
    add_para(doc, '')
    add_styled_table(doc, ['Livrabil', 'Detalii'], [
        ['Raport Audit PDF', 'Core Web Vitals, bottlenecks, recomandări'],
        ['Optimizări implementate', 'Cache, CDN, DB, imagini, CSS/JS pe ambele site-uri'],
        ['Raport before/after', 'Metrici comparative pre și post optimizare'],
        ['Testare stres', 'Simulare 150 utilizatori concurenți'],
    ])
    
    add_para(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Investiție Pachet 1: ________ EUR + TVA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    doc.add_page_break()
    
    # ── Pachet 2 ──
    add_heading_styled(doc, '3. Pachet 2 – Optimizare Checkout & Flux E-commerce', level=1)
    
    add_heading_styled(doc, '3.1 Simplificare checkout', level=2)
    add_bullet(doc, 'Audit UX al flow-ului actual CartFlow – identificare friction points')
    add_bullet(doc, 'Reducere câmpuri formular la minimum necesar')
    add_bullet(doc, 'Optimizare checkout mobil (autocomplete, input types corecte)')
    add_bullet(doc, 'Implementare skip cart → direct checkout unde e cazul')
    add_bullet(doc, 'Optimizare pagini upsell CartFlow (viteză + UX)')
    
    add_heading_styled(doc, '3.2 Configurare fluxuri automate e-commerce', level=2)
    add_bullet(doc, 'Add to Cart tracking (events pentru analytics)')
    add_bullet(doc, 'Email confirmare comandă (template optimizat)')
    add_bullet(doc, 'SMS automat cu informații preluare comandă')
    add_bullet(doc, 'SMS automat cu AWB tracking colet (integrare Cargus)')
    add_bullet(doc, 'Setup Abandoned Cart recovery – email sequence (2-3 emailuri)')
    add_bullet(doc, 'Setup Abandoned Cart recovery – SMS sequence')
    add_bullet(doc, 'Email-uri post-achiziție (follow-up, cerere review)')
    
    add_para(doc, '')
    add_styled_table(doc, ['Livrabil', 'Detalii'], [
        ['Checkout optimizat', 'Flow simplificat pe ambele magazine'],
        ['Abandoned Cart', 'Email + SMS sequence configurate'],
        ['Tracking AWB', 'SMS automat la expediere colet'],
        ['Post-achiziție', 'Email follow-up + review request'],
        ['Documentație', 'Flux complet documentat pentru echipa internă'],
    ])
    
    add_para(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Investiție Pachet 2: ________ EUR + TVA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    doc.add_page_break()
    
    # ── Pachet 3 ──
    add_heading_styled(doc, '4. Pachet 3 – Setup Automatizări WhatsApp', level=1)
    
    add_bullet(doc, 'Configurare canal WhatsApp Business API (prin provider: WATI, Twilio sau similar)')
    add_bullet(doc, 'Creare și aprobare template-uri mesaje automate:')
    add_bullet(doc, 'Confirmare comandă', level=1)
    add_bullet(doc, 'Reminder eveniment (cu 24h și 1h înainte)', level=1)
    add_bullet(doc, 'AWB tracking colet', level=1)
    add_bullet(doc, 'Follow-up post-eveniment cu oferta dedicată', level=1)
    add_bullet(doc, 'Integrare cu WooCommerce (trigger automat pe status comandă)')
    add_bullet(doc, 'Testare end-to-end pe ambele magazine')
    add_bullet(doc, 'Documentare workflow complet pentru echipa internă')
    
    add_para(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Notă: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('Costul lunar al provider-ului WhatsApp Business API (ex: WATI ~50€/lună) este suportat separat de client. Această ofertă acoperă setup-ul și configurarea completă.')
    run.font.size = Pt(10)
    
    add_para(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Investiție Pachet 3: ________ EUR + TVA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    doc.add_page_break()
    
    # ── Pachet 4 ──
    add_heading_styled(doc, '5. Pachet 4 – Optimizare Fluxuri Email (Newsman)', level=1)
    add_para(doc, 'Optimizare pe infrastructura existentă – fără schimbare platformă.', bold=True)
    
    add_bullet(doc, 'Audit segmente existente + recomandări segmentare avansată')
    add_bullet(doc, 'Creare/optimizare template-uri email:')
    add_bullet(doc, 'Reminder pre-eveniment (sequence 3 emailuri)', level=1)
    add_bullet(doc, 'Follow-up post-eveniment (sequence 2-3 emailuri)', level=1)
    add_bullet(doc, 'Abandoned cart (sequence 2 emailuri)', level=1)
    add_bullet(doc, 'Welcome sequence lead nou', level=1)
    add_bullet(doc, 'Optimizare subject lines + preview text (creștere open rate)')
    add_bullet(doc, 'Setup A/B testing pe campanii recurente')
    add_bullet(doc, 'Recomandări creștere CTR (butoane, layout, copy guidelines)')
    
    add_para(doc, '')
    add_styled_table(doc, ['Metric actual', 'Valoare', 'Obiectiv'], [
        ['Open Rate', '25% – 45%', 'Menținere/creștere cu segmentare'],
        ['Click Rate', '0.6% – 3%', 'Minim 3% consistent'],
        ['CTOR', '10% – 21%', 'Stabilizare peste 15%'],
    ])
    
    add_para(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Investiție Pachet 4: ________ EUR + TVA')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    doc.add_page_break()
    
    # ── Sumar ──
    add_heading_styled(doc, '6. Sumar Investiție & Condiții', level=1)
    
    add_styled_table(doc, ['Pachet', 'Descriere', 'Investiție'], [
        ['1. Performanță Tehnică', 'Audit + optimizare viteză + stabilitate peak traffic', '________ EUR'],
        ['2. Checkout & Flux E-com', 'Checkout simplificat + automatizări e-commerce', '________ EUR'],
        ['3. WhatsApp Setup', 'Configurare WhatsApp Business API + automatizări', '________ EUR'],
        ['4. Email Newsman', 'Optimizare segmente, template-uri, A/B testing', '________ EUR'],
        ['TOTAL', 'Toate cele 4 pachete', '________ EUR + TVA'],
    ])
    
    add_para(doc, '')
    add_heading_styled(doc, 'Condiții de plată', level=2)
    add_bullet(doc, '50% avans la semnarea contractului')
    add_bullet(doc, '50% la finalizarea și predarea proiectului')
    
    add_heading_styled(doc, 'Timeline estimat', level=2)
    add_bullet(doc, 'Audit: 3-5 zile lucrătoare')
    add_bullet(doc, 'Implementare: 10-15 zile lucrătoare')
    add_bullet(doc, 'Testare + handoff: 3 zile lucrătoare')
    add_bullet(doc, '', bold_prefix='Total: ~3-4 săptămâni')
    
    add_heading_styled(doc, 'Garanție post-livrare', level=2)
    add_bullet(doc, '30 de zile: fix gratuit pentru orice problemă cauzată de optimizările noastre')
    add_bullet(doc, 'Nu include probleme cauzate de update-uri WordPress/pluginuri făcute de echipa internă')
    
    add_heading_styled(doc, 'Ce NU include această ofertă', level=2)
    add_bullet(doc, 'Mentenanță lunară continuă')
    add_bullet(doc, 'Management campanii advertising (Facebook, TikTok, Google)')
    add_bullet(doc, 'Dezvoltare funcționalități noi / features custom')
    add_bullet(doc, 'Creare conținut (copywriting, design grafic)')
    add_bullet(doc, 'Suport tehnic ongoing post-livrare (după cele 30 zile garanție)')
    
    add_para(doc, '')
    add_heading_styled(doc, 'Livrabile finale', level=2)
    add_bullet(doc, 'Raport audit tehnic complet (PDF)')
    add_bullet(doc, 'Toate optimizările implementate pe ambele magazine')
    add_bullet(doc, 'Documentație tehnică completă (ce s-a modificat, cum se menține)')
    add_bullet(doc, 'Raport before/after (viteze, Core Web Vitals)')
    add_bullet(doc, 'Sesiune handoff 1h cu echipa IT RestartiX')
    
    # Save
    path = os.path.join(OUT_DIR, "Oferta-Optimizare-RestartiX.docx")
    doc.save(path)
    print(f"✅ Oferta salvata: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# DOCUMENT 2: GHID IMPLEMENTARE
# ═══════════════════════════════════════════════════════════

def create_ghid():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Title ──
    title = doc.add_heading('Ghid Implementare Pas cu Pas', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Optimizare WooCommerce – RestartiX')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_para(doc, 'Document intern – Full Out Media', size=10)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # FAZA 1
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'FAZA 1: Audit & Performanță Tehnică', level=1)
    add_para(doc, 'Durată estimată: 5-7 zile', bold=True)
    
    # Step 1.1
    add_heading_styled(doc, 'Pas 1.1 – Audit inițial', level=2)
    add_para(doc, 'Acces necesar:', bold=True)
    add_bullet(doc, 'Acces wp-admin pe ambele magazine (admin)')
    add_bullet(doc, 'Acces SSH/SFTP pe VPS (magazinrestartix.ro)')
    add_bullet(doc, 'Acces cPanel/hosting panel Romarg (zerodurere.net)')
    add_bullet(doc, 'Acces Cloudflare (dacă există) sau DNS management')
    
    add_para(doc, 'Ce facem:', bold=True)
    add_bullet(doc, 'Rulăm GTmetrix, PageSpeed Insights, WebPageTest pe ambele site-uri')
    add_bullet(doc, 'Salvăm screenshot-uri și scoruri BEFORE (le punem în raport)')
    add_bullet(doc, 'Verificăm versiuni: WordPress, PHP, MySQL, WooCommerce, CartFlow')
    add_bullet(doc, 'Listăm TOATE pluginurile active + inactive → identificăm pluginuri grele/redundante')
    add_bullet(doc, 'Verificăm wp_options → autoload bloat (SELECT SUM(LENGTH(option_value)) FROM wp_options WHERE autoload="yes")')
    add_bullet(doc, 'Verificăm cron jobs WordPress (wp-cron vs real cron)')
    add_bullet(doc, 'Pe VPS: verificăm PHP workers, memory_limit, max_execution_time, OPcache status')
    add_bullet(doc, 'Pe Romarg: verificăm limitările planului WordPress Pro-4')
    
    # Step 1.2
    add_heading_styled(doc, 'Pas 1.2 – Configurare cache', level=2)
    add_para(doc, 'VPS (magazinrestartix.ro):', bold=True)
    add_bullet(doc, 'Instalăm Redis server (apt install redis-server)')
    add_bullet(doc, 'Configurăm Redis Object Cache în WordPress (plugin Redis Object Cache)')
    add_bullet(doc, 'Instalăm plugin cache: WP Rocket sau LiteSpeed Cache (în funcție de web server)')
    add_bullet(doc, 'Configurăm page cache, browser cache, GZIP/Brotli compression')
    add_bullet(doc, 'Setăm cache exclusions: /cart, /checkout, /my-account, pagini CartFlow')
    
    add_para(doc, 'Romarg (zerodurere.net):', bold=True)
    add_bullet(doc, 'Verificăm ce cache oferă hostingul (LiteSpeed?)')
    add_bullet(doc, 'Configurăm LiteSpeed Cache plugin sau WP Super Cache')
    add_bullet(doc, 'Dacă Redis nu e disponibil: folosim object cache pe file-based')
    
    # Step 1.3
    add_heading_styled(doc, 'Pas 1.3 – Optimizare imagini', level=2)
    add_bullet(doc, 'Instalăm ShortPixel sau Imagify pe ambele site-uri')
    add_bullet(doc, 'Bulk optimize toate imaginile existente → WebP')
    add_bullet(doc, 'Configurăm lazy loading nativ WordPress (loading="lazy")')
    add_bullet(doc, 'Verificăm dacă tema face preload la LCP image (hero)')
    add_bullet(doc, 'Adăugăm fetchpriority="high" pe imaginea principală')
    
    # Step 1.4
    add_heading_styled(doc, 'Pas 1.4 – Optimizare CSS/JS', level=2)
    add_bullet(doc, 'Cu WP Rocket: activăm Minify CSS, Minify JS, Combine JS (cu atenție)')
    add_bullet(doc, 'Delay JS execution pentru: analytics, chat widgets, Facebook Pixel, TikTok Pixel')
    add_bullet(doc, 'Remove unused CSS per pagină (dacă WP Rocket Pro, sau cu Perfmatters)')
    add_bullet(doc, 'Preload fonturi Google Fonts (sau self-host)')
    add_bullet(doc, 'Eliminăm pluginuri care încarcă CSS/JS global dar sunt folosite pe 1-2 pagini')
    
    # Step 1.5
    add_heading_styled(doc, 'Pas 1.5 – Optimizare bază de date', level=2)
    add_bullet(doc, 'Cleanup cu WP-Optimize sau Advanced Database Cleaner:')
    add_bullet(doc, 'Ștergem post revisions (păstrăm ultimele 3)', level=1)
    add_bullet(doc, 'Ștergem auto-drafts, trash posts', level=1)
    add_bullet(doc, 'Ștergem transients expirate', level=1)
    add_bullet(doc, 'Ștergem WooCommerce expired sessions', level=1)
    add_bullet(doc, 'Optimizăm tabele MySQL (OPTIMIZE TABLE)', level=1)
    add_bullet(doc, 'Adăugăm în wp-config.php: define("WP_POST_REVISIONS", 3);')
    add_bullet(doc, 'Dezactivăm wp-cron.php și setăm real cron pe VPS (crontab -e → */5 * * * * curl ...)')
    
    # Step 1.6
    add_heading_styled(doc, 'Pas 1.6 – CDN & stabilitate peak traffic', level=2)
    add_bullet(doc, 'Setup Cloudflare (free plan e suficient):')
    add_bullet(doc, 'Configurăm DNS nameservers → Cloudflare', level=1)
    add_bullet(doc, 'Page Rules: cache everything pe /wp-content/uploads/*', level=1)
    add_bullet(doc, 'Bypass cache pe /cart/*, /checkout/*, /my-account/*', level=1)
    add_bullet(doc, 'Activăm Brotli compression, Auto Minify, Rocket Loader OFF', level=1)
    add_bullet(doc, 'Bot Fight Mode activat', level=1)
    add_bullet(doc, 'Rate limiting (optional, Cloudflare Pro) → max 30 req/min per IP pe checkout', level=1)
    add_bullet(doc, 'Pe VPS: tuning PHP-FPM:')
    add_bullet(doc, 'pm = dynamic, pm.max_children = 30-50 (în funcție de RAM)', level=1)
    add_bullet(doc, 'pm.start_servers = 10, pm.min_spare = 5, pm.max_spare = 20', level=1)
    add_bullet(doc, 'OPcache: opcache.memory_consumption=256, opcache.max_accelerated_files=20000', level=1)
    
    # Step 1.7
    add_heading_styled(doc, 'Pas 1.7 – Testare stres', level=2)
    add_bullet(doc, 'Folosim k6.io (gratuit) sau loader.io pentru load testing')
    add_bullet(doc, 'Simulăm 150 utilizatori concurenți pe homepage + pe o pagină de produs + checkout')
    add_bullet(doc, 'Verificăm TTFB sub 500ms, error rate 0%')
    add_bullet(doc, 'Documentăm rezultatele în raportul final')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # FAZA 2
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'FAZA 2: Checkout & Automatizări E-commerce', level=1)
    add_para(doc, 'Durată estimată: 5-7 zile', bold=True)
    
    add_heading_styled(doc, 'Pas 2.1 – Audit checkout actual', level=2)
    add_bullet(doc, 'Parcurgem manual tot flow-ul de cumpărare pe ambele magazine (desktop + mobil)')
    add_bullet(doc, 'Identificăm: câte click-uri până la finalizare, câmpuri inutile, erori UX')
    add_bullet(doc, 'Verificăm CartFlow: câte steps are, cum arată upsell-urile')
    add_bullet(doc, 'Testăm pe mobil (60%+ din trafic vine de pe mobil la evenimente)')
    
    add_heading_styled(doc, 'Pas 2.2 – Simplificare checkout', level=2)
    add_bullet(doc, 'Reducem câmpurile formular:')
    add_bullet(doc, 'Eliminăm "Company" dacă nu e necesar', level=1)
    add_bullet(doc, 'Adăugăm autocomplete pe adresă', level=1)
    add_bullet(doc, 'Input type="tel" pe telefon, type="email" pe email', level=1)
    add_bullet(doc, 'Pre-fill date dacă clientul e logat', level=1)
    add_bullet(doc, 'Adăugăm trust badges lângă butonul de plată (EuPlatesc logo, SSL badge)')
    add_bullet(doc, 'Dacă produsul e digital → ascundem câmpurile de livrare complet')
    add_bullet(doc, 'Verificăm dacă CartFlow bypass cart funcționează (link produs → direct checkout)')
    
    add_heading_styled(doc, 'Pas 2.3 – Abandoned Cart', level=2)
    add_bullet(doc, 'Instalăm plugin: CartFlows Pro (dacă nu au deja) sau WooCommerce Cart Abandonment Recovery (gratuit)')
    add_bullet(doc, 'Configurăm sequence:')
    add_bullet(doc, 'Email 1: la 1 oră după abandon – "Ai uitat ceva în coș"', level=1)
    add_bullet(doc, 'Email 2: la 24 ore – reminder cu urgență', level=1)
    add_bullet(doc, 'Email 3: la 72 ore – ultimul reminder (opțional: mic discount)', level=1)
    add_bullet(doc, 'SMS (prin Newsman sau provider SMS):')
    add_bullet(doc, 'SMS 1: la 2 ore după abandon', level=1)
    add_bullet(doc, 'Includem link direct la checkout cu produsele în coș', level=1)
    
    add_heading_styled(doc, 'Pas 2.4 – Tracking AWB & comunicare comandă', level=2)
    add_bullet(doc, 'Verificăm integrarea Cargus → extrage AWB automat la generare')
    add_bullet(doc, 'Configurăm email WooCommerce pe status "Shipped/Completed" cu AWB tracking link')
    add_bullet(doc, 'Adăugăm SMS automat la expediere (prin Newsman SMS sau plugin dedicat)')
    add_bullet(doc, 'Template SMS: "Comanda ta #{order_id} a fost expediată. Tracking: {awb_link}"')
    
    add_heading_styled(doc, 'Pas 2.5 – Email-uri post-achiziție', level=2)
    add_bullet(doc, 'Email follow-up la 7 zile: "Cum a fost experiența?"')
    add_bullet(doc, 'Email cerere review la 14 zile (pentru produse fizice)')
    add_bullet(doc, 'Pentru produse digitale: email la 3 zile "Ai reușit să accesezi platforma?"')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # FAZA 3
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'FAZA 3: Setup WhatsApp Business', level=1)
    add_para(doc, 'Durată estimată: 3-5 zile', bold=True)
    
    add_heading_styled(doc, 'Pas 3.1 – Alegere provider', level=2)
    add_para(doc, 'Opțiuni recomandate:', bold=True)
    add_styled_table(doc, ['Provider', 'Preț/lună', 'Avantaje'], [
        ['WATI.io', '~49€', 'Cel mai ușor de configurat, UI bun, integrare WooCommerce directă'],
        ['Twilio', 'Pay per message', 'Flexibil, API puternic, dar necesită dezvoltare custom'],
        ['360dialog', '~50€', 'Official WhatsApp partner, bun pentru volume mari'],
    ])
    add_para(doc, '')
    add_para(doc, 'Recomandare: WATI.io – cel mai rapid setup, are plugin WooCommerce ready-made.', bold=True)
    
    add_heading_styled(doc, 'Pas 3.2 – Setup WhatsApp Business API', level=2)
    add_bullet(doc, 'Creăm cont pe WATI.io cu numărul de telefon business RestartiX')
    add_bullet(doc, 'Verificăm Facebook Business Manager (necesar pentru WhatsApp API)')
    add_bullet(doc, 'Conectăm numărul la WATI → verificare prin cod SMS')
    add_bullet(doc, 'Instalăm plugin WATI pe ambele magazine WooCommerce')
    
    add_heading_styled(doc, 'Pas 3.3 – Creare template-uri mesaje', level=2)
    add_para(doc, 'IMPORTANT: Template-urile WhatsApp trebuie aprobate de Meta (24-48h).', bold=True)
    add_bullet(doc, 'Template 1 – Confirmare comandă:')
    add_bullet(doc, '"Salut {{name}}! 🎉 Comanda ta #{{order_id}} a fost înregistrată. Te vom ține la curent cu statusul."', level=1)
    add_bullet(doc, 'Template 2 – Reminder eveniment 24h:')
    add_bullet(doc, '"Salut {{name}}! 📅 Mâine la ora {{time}} te așteptăm la {{event_name}}. Link acces: {{link}}"', level=1)
    add_bullet(doc, 'Template 3 – Reminder eveniment 1h:')
    add_bullet(doc, '"{{name}}, evenimentul {{event_name}} începe în curând! 🔴 Intră acum: {{link}}"', level=1)
    add_bullet(doc, 'Template 4 – AWB Tracking:')
    add_bullet(doc, '"Comanda ta #{{order_id}} a fost expediată! 🚚 Tracking: {{tracking_link}}"', level=1)
    add_bullet(doc, 'Template 5 – Follow-up post-eveniment:')
    add_bullet(doc, '"Salut {{name}}! Mulțumim că ai participat la {{event_name}}. Oferta specială e disponibilă aici: {{link}}"', level=1)
    
    add_heading_styled(doc, 'Pas 3.4 – Configurare triggere automate', level=2)
    add_bullet(doc, 'În WATI / plugin WooCommerce:')
    add_bullet(doc, 'Status "Processing" → trimite Template 1 (confirmare)', level=1)
    add_bullet(doc, 'Status "Completed/Shipped" → trimite Template 4 (AWB)', level=1)
    add_bullet(doc, 'Reminder-ele de eveniment → se configurează manual sau prin Zapier/Make.com din calendarul de evenimente', level=1)
    add_bullet(doc, 'Follow-up post-eveniment → trigger la 2h după ora de final eveniment', level=1)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # FAZA 4
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'FAZA 4: Optimizare Email Marketing (Newsman)', level=1)
    add_para(doc, 'Durată estimată: 3-5 zile', bold=True)
    
    add_heading_styled(doc, 'Pas 4.1 – Audit segmente', level=2)
    add_bullet(doc, 'Exportăm lista segmentelor existente din Newsman')
    add_bullet(doc, 'Analizăm: câte segmente active, criterii de segmentare, overlap')
    add_bullet(doc, 'Recomandări segmente noi:')
    add_bullet(doc, 'Clienți activi (comandă în ultimele 90 zile)', level=1)
    add_bullet(doc, 'Clienți inactivi (fără comandă > 90 zile)', level=1)
    add_bullet(doc, 'Lead-uri reci (înscriși la eveniment, fără achiziție)', level=1)
    add_bullet(doc, 'VIP (>2 comenzi sau valoare totală >X RON)', level=1)
    add_bullet(doc, 'Per tip produs (fizic vs digital)', level=1)
    
    add_heading_styled(doc, 'Pas 4.2 – Optimizare template-uri', level=2)
    add_bullet(doc, 'Redesign template-uri cu focus pe:')
    add_bullet(doc, 'Un singur CTA clar per email (creștere CTR)', level=1)
    add_bullet(doc, 'Butoane mari, vizibile, culoare contrastantă', level=1)
    add_bullet(doc, 'Preview text optimizat (nu repetă subject line)', level=1)
    add_bullet(doc, 'Layout mobile-first (65%+ deschideri pe mobil)', level=1)
    add_bullet(doc, 'Personalizare cu {{first_name}} în subject + body', level=1)
    
    add_heading_styled(doc, 'Pas 4.3 – Setup sequences', level=2)
    add_para(doc, 'Reminder pre-eveniment (3 emailuri):', bold=True)
    add_bullet(doc, 'E1: Cu 3 zile înainte – "Te-ai înscris! Iată ce te așteaptă"')
    add_bullet(doc, 'E2: Cu 1 zi înainte – "Mâine e ziua cea mare"')
    add_bullet(doc, 'E3: Cu 1h înainte – "Începem în curând! Link acces"')
    
    add_para(doc, 'Follow-up post-eveniment (3 emailuri):', bold=True)
    add_bullet(doc, 'E1: La 2h după – "Mulțumim! Iată oferta specială" (cu deadline)')
    add_bullet(doc, 'E2: La 24h – "Oferta expiră în curând" (urgență)')
    add_bullet(doc, 'E3: La 48h – "Ultima șansă" (doar pentru cei care n-au cumpărat)')
    
    add_para(doc, 'Welcome sequence lead nou (2 emailuri):', bold=True)
    add_bullet(doc, 'E1: Imediat – "Bun venit! Cine suntem și ce facem"')
    add_bullet(doc, 'E2: La 3 zile – "Următorul eveniment gratuit"')
    
    add_heading_styled(doc, 'Pas 4.4 – A/B Testing', level=2)
    add_bullet(doc, 'Configurăm A/B test pe subject line pentru fiecare campanie majoră')
    add_bullet(doc, 'Split: 20% audiență test (10%+10%), 80% primesc varianta câștigătoare')
    add_bullet(doc, 'Testăm: emojis vs fără, întrebare vs afirmație, lungime scurtă vs lungă')
    add_bullet(doc, 'Documentăm rezultatele într-un spreadsheet de learnings')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # FAZA 5
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'FAZA 5: Testare Finală & Handoff', level=1)
    add_para(doc, 'Durată estimată: 2-3 zile', bold=True)
    
    add_heading_styled(doc, 'Pas 5.1 – Testare completă', level=2)
    add_bullet(doc, 'Testăm tot flow-ul end-to-end pe ambele magazine:')
    add_bullet(doc, 'Adaugă produs → Cart → Checkout → Plată test EuPlatesc → Confirmare', level=1)
    add_bullet(doc, 'Verificăm: email confirmare, SMS, WhatsApp, MiniCRM sync', level=1)
    add_bullet(doc, 'Testăm abandoned cart: adăugăm în coș, abandonăm, verificăm emailurile', level=1)
    add_bullet(doc, 'Testăm pe mobil (iOS + Android)', level=1)
    add_bullet(doc, 'Re-testăm viteza: GTmetrix, PageSpeed → screenshot AFTER', level=1)
    
    add_heading_styled(doc, 'Pas 5.2 – Documentație', level=2)
    add_bullet(doc, 'Documentăm tot ce s-a modificat:')
    add_bullet(doc, 'Pluginuri instalate/dezinstalate', level=1)
    add_bullet(doc, 'Configurări server (PHP, MySQL, Redis, Cloudflare)', level=1)
    add_bullet(doc, 'Fluxuri email/SMS/WhatsApp cu diagrame', level=1)
    add_bullet(doc, 'Credențiale noi create (WATI, Cloudflare etc.)', level=1)
    
    add_heading_styled(doc, 'Pas 5.3 – Sesiune handoff', level=2)
    add_bullet(doc, 'Call 1h cu echipa IT RestartiX')
    add_bullet(doc, 'Prezentăm: ce s-a făcut, cum funcționează, ce trebuie menținut')
    add_bullet(doc, 'Predăm documentația + acces la toate tool-urile configurate')
    add_bullet(doc, 'Stabilim contact 30 zile garanție')
    
    # Save
    path = os.path.join(OUT_DIR, "Ghid-Implementare-RestartiX.docx")
    doc.save(path)
    print(f"✅ Ghid salvat: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# RUN
# ═══════════════════════════════════════════════════════════

if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    create_oferta()
    create_ghid()
    print("\n🎉 Ambele documente generate cu succes!")
